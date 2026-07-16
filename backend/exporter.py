import io
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor

def export_docx(content: str, title: str) -> io.BytesIO:
    """Generate a clean Word document for the draft."""
    doc = Document()
    # Add Title
    doc.add_heading(title, level=1)
    
    # Add spacing and body
    for block in content.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("Subject:"):
            p = doc.add_paragraph()
            run = p.add_run(block)
            run.bold = True
        elif block.startswith("DRAFT:") or block.isupper():
            doc.add_heading(block, level=2)
        else:
            doc.add_paragraph(block)
            
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def export_pdf(content: str, title: str) -> io.BytesIO:
    """Generate a premium styled PDF document for the draft."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    # Custom premium styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=HexColor('#004ac6'),
        spaceAfter=15
    )
    
    subject_style = ParagraphStyle(
        'DocSubject',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=16,
        textColor=HexColor('#191c1d'),
        spaceAfter=12
    )
    
    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=15,
        textColor=HexColor('#2e3132'),
        spaceAfter=10
    )
    
    meta_style = ParagraphStyle(
        'DocMeta',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=8,
        leading=10,
        textColor=HexColor('#737686'),
        spaceAfter=20
    )
    
    story = []
    
    # Header tag
    story.append(Paragraph("NYAYSETU PREC_ADVOCATE AI &bull; CONFIDENTIAL LEGAL DRAFT", meta_style))
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 10))
    
    for block in content.split("\n\n"):
        block = block.strip().replace("\n", "<br/>")
        if not block:
            continue
        if block.startswith("Subject:"):
            story.append(Paragraph(block, subject_style))
        else:
            story.append(Paragraph(block, body_style))
            
    doc.build(story)
    buffer.seek(0)
    return buffer
